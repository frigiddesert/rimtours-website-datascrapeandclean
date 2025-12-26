import json
import os
import re
import pandas as pd
import docx
from difflib import SequenceMatcher

# ==========================================
# CONFIGURATION
# ==========================================
OUTPUT_DIR = "outline_final_import"
TOURS_FILE = "unified_tours.json"
DOCS_FOLDER = "itinerary_docs"

# ==========================================
# 1. HELPERS
# ==========================================
def extract_stats_from_text(text):
    """Extract Miles, Elev, Camp from raw string."""
    if not text: return {"Miles": "", "Elev": "", "Camp": ""}
    
    stats = {"Miles": "", "Elev": "", "Camp": ""}
    # Miles
    m = re.search(r"(\d+(?:\.\d+)?)\s*(?:miles|mi\b)", text, re.IGNORECASE)
    if m: stats["Miles"] = m.group(1)
    # Elev
    e = re.search(r"(\d{1,3}(?:,\d{3})?)\s*(?:ft|feet|’|')\s*(?:gain|climb)", text, re.IGNORECASE)
    if e: stats["Elev"] = e.group(1)
    # Camp
    c = re.search(r"(?:Camp|Lodging|Overnight):\s*(.*)", text, re.IGNORECASE)
    if c: stats["Camp"] = c.group(1).strip()
    return stats

def load_docx_itineraries():
    """Load all itinerary documents."""
    db = {}
    if not os.path.exists(DOCS_FOLDER): return db
    
    files = [f for f in os.listdir(DOCS_FOLDER) if f.lower().endswith(('.doc', '.docx'))]
    print(f"📁 Found {len(files)} itinerary docs...")
    
    for f in files:
        filepath = os.path.join(DOCS_FOLDER, f)
        try:
            if f.lower().endswith('.docx'):
                doc = docx.Document(filepath)
                db[f] = doc
            elif f.lower().endswith('.doc'):
                import subprocess
                result = subprocess.run(['antiword', filepath], 
                                      capture_output=True, check=True, text=True)
                text = result.stdout
                # Create mock object with paragraphs for consistency
                class MockDoc:
                    def __init__(self, text):
                        self.paragraphs = [MockPara(line) for line in text.split('\n') if line.strip()]
                class MockPara:
                    def __init__(self, text):
                        self.text = text
                db[f] = MockDoc(text)
        except Exception as e:
            print(f"  Error reading {f}: {e}")
            continue
    return db

def process_doc_content(doc):
    """Parse loaded Docx/Doc object into our list format."""
    days = []
    day_pattern = re.compile(r'(?:^|\n)\s*(?:Day|DAY|Stage|STAGE)\s+(\d+|[Oo]ne|[Tt]wo|[Tt]hree|[Ff]our|[Ff]ive|[Ss]ix)(?:[:\.-]|\s|$)', re.IGNORECASE)
    
    # Get all text first
    full_text = '\n'.join([para.text for para in doc.paragraphs if para.text])
    
    # Find all day markers
    matches = list(day_pattern.finditer(full_text))
    
    if not matches:
        # If no day markers, return as single day
        stats = extract_stats_from_text(full_text)
        return [{
            "Day": "1",
            "Content": full_text[:500],
            "Miles": stats['Miles'],
            "Elev": stats['Elev'],
            "Camp": stats['Camp']
        }]
    
    # Process each day segment
    mapping = {'one':'1','two':'2','three':'3','four':'4','five':'5','six':'6'}
    
    for i in range(len(matches)):
        match = matches[i]
        day_num = mapping.get(match.group(1).lower(), match.group(1))
        
        # Determine content boundaries
        start_pos = match.end()
        end_pos = matches[i+1].start() if i+1 < len(matches) else len(full_text)
        
        content = full_text[start_pos:end_pos].strip()
        
        # Extract stats from content
        stats = extract_stats_from_text(content)
        
        days.append({
            "Day": day_num,
            "Content": content,
            "Miles": stats['Miles'],
            "Elev": stats['Elev'],
            "Camp": stats['Camp']
        })
    
    return days

def categorize(tour):
    name = tour['Master_Name'].lower()
    variants = tour['Arctic_Variants']
    if "oregon" in name or "sierra" in name: return "Archive"
    if "rental" in name or "service" in name or "shuttle" in name: return "Rentals & Services"
    is_multi = False
    for v in variants:
        bg = str(v.get('Business_Group', ''))
        if bg in ['3','4','23','26']: is_multi = True
    if not variants and ("4-day" in name or "3-day" in name): is_multi = True
    
    if not is_multi: return "Day Tours"
    if "arizona" in name: return "Multi-Day Tours/Arizona"
    if "colorado" in name: return "Multi-Day Tours/Colorado"
    return "Multi-Day Tours/Utah"

def generate_markdown(tour, itinerary):
    title = tour['Master_Name']
    meta = tour.get('Meta', {})
    
    # 1. Get Primary Arctic Shortname (from the first variant, or list them)
    shortnames = sorted(list(set([v['Shortname'] for v in tour['Arctic_Variants'] if v.get('Shortname')])))
    arctic_code = ", ".join(shortnames) if shortnames else "N/A"

    # 2. Variants Table
    v_rows = []
    for v in tour['Arctic_Variants']:
        # Use the actual field names from the data
        aid = str(v.get('id', 'N/A'))
        v_name = v.get('name', 'N/A')
        v_duration = v.get('duration', 'N/A')
        v_type = v.get('name', 'N/A')  # or extract from keywords

        # Get price from v if available, otherwise use TBD
        v_price = v.get('Price', 'TBD')
        v_rows.append(f"| **{v_name}** | {aid} | {v_price} | {v_duration} | {v_type} |")
    v_table = "\n".join(v_rows) if v_rows else "| No Linked Variants | - | - | - | - |"

    # 3. Itinerary Table
    i_rows = []
    full_text = ""
    if itinerary:
        for d in itinerary:
            desc = d['Content'][:100].replace("\n"," ") + "..."
            i_rows.append(f"| {d['Day']} | {desc} | {d['Miles']} | {d['Elev']} | {d['Camp']} | All |")
            full_text += f"### Day {d['Day']}\n{d['Content']}\n\n"
    else:
        i_rows = ["| 1 | | | | | All |"]
    i_table = "\n".join(i_rows)

    # 4. Fees Block - Extract from Meta
    fees_section = ""
    # Extract actual fee values from the metadata (they might be in the long description or other fields)
    if tour.get('Description_Long'):
        # Look for pricing info in the tour description
        desc = tour['Description_Long']
        # Look for patterns like "$XX fee" or "X/day rental"
        bike_fee_match = re.search(r'(?:Bike|Bicycle).*?\$(\d+)', desc, re.IGNORECASE)
        camp_fee_match = re.search(r'(?:Camp|Kit).*?\$(\d+)', desc, re.IGNORECASE)
        shuttle_fee_match = re.search(r'(?:Shuttle|Transport).*?\$(\d+)', desc, re.IGNORECASE)
        
        bike_fee = f"${bike_fee_match.group(1)}/day" if bike_fee_match else "N/A"
        camp_fee = f"${camp_fee_match.group(1)}/kit" if camp_fee_match else "N/A"
        shuttle_fee = f"${shuttle_fee_match.group(1)}/trip" if shuttle_fee_match else "Included or N/A"
        
        fees_section = f"""
## 💰 Fees & Logistics
| Item | Cost / Details |
| :--- | :--- |
| **Bike Rental** | {bike_fee} |
| **Camp Kit** | {camp_fee} |
| **Shuttle Service** | {shuttle_fee} |
"""
    elif any(meta.get(field) for field in ['Fees_Bike', 'Fees_Camp', 'Fees_Shuttle']):
        fees_section = f"""
## 💰 Fees & Logistics
| Item | Cost / Details |
| :--- | :--- |
| **Bike Rental** | {meta.get('Fees_Bike', 'N/A')} |
| **Camp Kit** | {meta.get('Fees_Camp', 'N/A')} |
| **Shuttle Service** | {meta.get('Fees_Shuttle', 'Included or N/A')} |
"""

    return f"""# {title}

<!-- SYSTEM METADATA -->
| Arctic Code | System Status | Website ID |
| :--- | :--- | :--- |
| **{arctic_code}** | {tour.get('Sync_Status', 'New')} | {tour.get('Website_ID', 'N/A')} |

---

## 1. The Shared DNA
**Subtitle:** {meta.get('Subtitle', '')}  
**Region:** {meta.get('Region', '')}  
**Skill Level:** {meta.get('Skill', '')}  
**Season:** {meta.get('Season', '')}

**Short Description:**
> {tour.get('Description_Short', '')}

**Long Description:**
> {tour.get('Description_Long', '')[:1500]}...

**Images (Filenames):**
`{meta.get('Images', 'No images found')}`

---

## 2. Arctic Configurations (SKUs)
| Code | Variant Name | Arctic ID | Price | Duration |
| :--- | :--- | :--- | :--- | :--- |
{v_table}

{fees_section}

---

## 3. Itinerary Logic
| Day | Route Description | Miles | Elev Gain | Camp/Lodging | Applies To |
| :--- | :--- | :--- | :--- | :--- | :--- |
{i_table}

---

## 4. Full Itinerary Text
{full_text}
"""

def main():
    # Load Data
    try:
        with open(TOURS_FILE, 'r') as f: tours = json.load(f)
    except: return print("Run unify_tours_advanced.py first!")

    # Load Itineraries
    doc_db = load_docx_itineraries()

    # Create Output Dir
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)

    for tour in tours:
        name = tour['Master_Name']
        if any(x in name.lower() for x in ['fake', 'test', 'sold out', 'aaa ']): continue

        # 1. Find Itinerary Doc
        matched_doc_name = None
        itinerary_data = None

        # Fuzzy Search
        best_score = 0
        for fname in doc_db:
            clean = fname.replace(".docx","").replace(".doc","").replace("-"," ").replace("_"," ")
            score = SequenceMatcher(None, name.lower(), clean.lower()).ratio()
            if score > 0.65 and score > best_score:
                best_score = score
                matched_doc_name = fname

        if matched_doc_name:
            itinerary_data = process_doc_content(doc_db[matched_doc_name])
            print(f"✅ Found Doc for {name}: {matched_doc_name}")
        else:
            category = categorize(tour)
            if "Day Tours" in category:
                # Day tours don't need itineraries
                itinerary_data = None
                print(f"ℹ️  Day Tour (no itinerary needed): {name}")
            else:
                # Multi-day tours without docs
                print(f"⚠️  No Doc found for {name}")
                itinerary_data = None

        # Generate
        md = generate_markdown(tour, itinerary_data)

        # Save to proper subfolder
        sub = categorize(tour)
        p = os.path.join(OUTPUT_DIR, sub)
        if not os.path.exists(p): os.makedirs(p)
        safe = re.sub(r'[^\w\-_\. ]', '_', name) + ".md"
        with open(os.path.join(p, safe), 'w', encoding='utf-8') as f: f.write(md)

        print(f"📝 Generated {name}")

if __name__ == "__main__":
    main()