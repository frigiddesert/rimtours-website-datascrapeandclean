import json
import os
import re
import subprocess
from difflib import SequenceMatcher
import pandas as pd

# ==========================================
# CONFIGURATION
# ==========================================
OUTPUT_DIR = "outline_final_import"
TOURS_FILE = "unified_tours.json"
DOCS_FOLDER = "itinerary_docs"

# ==========================================
# 1. HELPERS
# ==========================================
def extract_stats_from_text(text):
    """Extract Miles, Elev, Camp from raw string."""
    if not text: return {"Miles": "", "Elev": "", "Camp": ""}
    
    stats = {"Miles": "", "Elev": "", "Camp": ""}
    # Miles
    m = re.search(r"(\d+(?:\.\d+)?)\s*(?:miles|mi\b)", text, re.IGNORECASE)
    if m: stats["Miles"] = m.group(1)
    # Elev
    e = re.search(r"(\d{1,3}(?:,\d{3})?)\s*(?:ft|feet|’|')\s*(?:gain|climb)", text, re.IGNORECASE)
    if e: stats["Elev"] = e.group(1)
    # Camp
    c = re.search(r"(?:Camp|Lodging|Overnight):\s*(.*)", text, re.IGNORECASE)
    if c: stats["Camp"] = c.group(1).strip()
    return stats

def load_docx_itineraries(directory=None):
    """Load all itinerary documents from specified directory or default location."""
    db = {}

    # Use provided directory or default to DOCS_FOLDER
    folder = directory if directory else DOCS_FOLDER

    if not os.path.exists(folder): return db

    files = [f for f in os.listdir(folder) if f.lower().endswith(('.doc', '.docx'))]
    print(f"📄 Found {len(files)} itinerary documents in {folder}...")

    for f in files:
        filepath = os.path.join(folder, f)
        try:
            if f.lower().endswith('.docx'):
                # Use docx for .docx files
                import docx
                doc = docx.Document(filepath)
                db[f] = doc
            elif f.lower().endswith('.doc'):
                # Use antiword for .doc files
                result = subprocess.run(['antiword', filepath],
                                      capture_output=True, text=True, check=True)
                text = result.stdout
                # Create a mock object with a paragraphs-like structure for consistency
                class MockDoc:
                    def __init__(self, text):
                        self.content = text
                db[f] = MockDoc(text)
        except Exception as e:
            print(f"  Error reading {f}: {e}")
            continue
    return db

def process_doc_content(doc_obj):
    """
    Parse loaded Docx/Doc object into our list format.
    Handles documents with two sections (summary + details) by prioritizing the more complete section.
    """
    if hasattr(doc_obj, 'paragraphs') and doc_obj.paragraphs:
        # This is a .docx file, get all text
        full_text = '\n'.join([para.text for para in doc_obj.paragraphs if para.text])
    else:
        # This is a .doc file (Antiword result stored in content attribute)
        full_text = doc_obj.content

    # Find all day markers
    day_pattern = re.compile(r'(?:^|\n)\s*(?:Day|DAY|Stage|STAGE)\s+(\d+|[Oo]ne|[Tt]wo|[Tt]hree|[Ff]our|[Ff]ive|[Ss]ix)(?:[:\.-]|\s|$)', re.IGNORECASE)
    
    matches = list(day_pattern.finditer(full_text))
    
    if not matches:
        # No day markers found, return as single day
        stats = extract_stats_from_text(full_text)
        return [{
            "Day": "1",
            "Content": full_text[:500],
            "Miles": stats['Miles'],
            "Elev": stats['Elev'],
            "Camp": stats['Camp']
        }]
    
    # Group matches by day number to handle duplicate sections
    mapping = {'one':'1','two':'2','three':'3','four':'4','five':'5','six':'6'}
    day_groups = {}
    
    for match in matches:
        day_num = mapping.get(match.group(1).lower(), match.group(1))
        if day_num not in day_groups:
            day_groups[day_num] = []
        day_groups[day_num].append(match)
    
    # For each day number, find the longest content segment
    final_days = []
    processed_ranges = set()
    
    for day_num, day_matches in day_groups.items():
        best_content = ""
        best_match = None
        
        for match in day_matches:
            # Determine content boundaries
            start_pos = match.end()
            # Find the next day marker or end of document
            next_starts = [next_m.start() for next_m in matches if next_m.start() > match.end()]
            end_pos = min(next_starts) if next_starts else len(full_text)
            
            # Avoid overlapping with already processed ranges
            content = full_text[start_pos:end_pos].strip()
            
            # Check overlap with already processed sections
            overlap = False
            for processed_start, processed_end in processed_ranges:
                if start_pos < processed_end and end_pos > processed_start:
                    overlap = True
                    break
            
            if not overlap and len(content) > len(best_content):
                best_content = content
                best_match = match
        
        if best_match:
            # Mark this range as processed to avoid duplication
            range_start = best_match.end()
            next_starts = [next_m.start() for next_m in matches if next_m.start() > best_match.end()]
            range_end = min(next_starts) if next_starts else len(full_text)
            processed_ranges.add((range_start, range_end))
            
            # Extract stats from the best content
            stats = extract_stats_from_text(best_content)
            
            final_days.append({
                "Day": day_num,
                "Content": best_content,
                "Miles": stats['Miles'],
                "Elev": stats['Elev'],
                "Camp": stats['Camp']
            })
    
    # Sort by day number to ensure proper ordering
    final_days.sort(key=lambda x: int(x['Day']) if x['Day'].isdigit() else x['Day'])
    return final_days

def categorize(tour):
    name = tour['Master_Name'].lower()
    variants = tour['Arctic_Variants']
    if "oregon" in name or "sierra" in name: return "Archive"
    if "rental" in name or "service" in name or "shuttle" in name: return "Rentals & Services"
    is_multi = False
    for v in variants:
        bg = str(v.get('Business_Group', ''))  # Correct field name
        if bg in ['3','4','23','26']: is_multi = True
    if not variants and ("4-day" in name or "3-day" in name): is_multi = True

    if not is_multi: return "Day Tours"
    if "arizona" in name: return "Multi-Day Tours/Arizona"
    if "colorado" in name: return "Multi-Day Tours/Colorado"
    return "Multi-Day Tours/Utah"

def generate_markdown(tour, itinerary):
    title = tour['Master_Name']
    meta = tour.get('Meta', {})
    
    # 1. Get Primary Arctic Shortname (from first variant or list all)
    shortnames = sorted(list(set([v.get('Shortname', '') for v in tour['Arctic_Variants'] if v.get('Shortname')])))
    arctic_code = ", ".join([sn for sn in shortnames if sn]) if shortnames else "N/A"

    # 2. Variants Table
    v_rows = []
    for v in tour['Arctic_Variants']:
        v_name = v.get('Name', 'N/A')
        aid = v.get('Arctic_ID', 'N/A')
        price = v.get('Price', 'TBD')
        duration = v.get('Duration', 'N/A')
        vtype = v.get('Type', 'N/A')
        
        v_rows.append(f"| **{v_name}** | {aid} | {price} | {duration} | {vtype} |")
    v_table = "\n".join(v_rows) if v_rows else "| No Linked Variants | - | - | - | - |"

    # 3. Itinerary Table
    i_rows = []
    full_text = ""
    if itinerary:
        for d in itinerary:
            desc = d['Content'][:100].replace("\n"," ") + "..."
            i_rows.append(f"| {d['Day']} | {desc} | {d['Miles']} | {d['Elev']} | {d['Camp']} | All |")
            full_text += f"### Day {d['Day']}\n{d['Content']}\n\n"
    else:
        i_rows = ["| 1 | | | | | All |"]
    i_table = "\n".join(i_rows)

    # 4. Fees Section (try to extract from meta or description)
    fees_section = ""
    # Extract fees information from meta fields
    bike_fee = meta.get('Fees_Bike', 'N/A')
    camp_fee = meta.get('Fees_Camp', 'N/A') 
    shuttle_fee = meta.get('Fees_Shuttle', 'N/A')
    
    # Also try to extract from long description if not in meta
    if bike_fee == 'N/A' and tour.get('Description_Long'):
        desc = tour['Description_Long']
        bike_match = re.search(r'Bike.*?\$(\d+)', desc, re.IGNORECASE)
        if bike_match: bike_fee = f"${bike_match.group(1)}/day"
        camp_match = re.search(r'Camp.*?\$(\d+)', desc, re.IGNORECASE)
        if camp_match: camp_fee = f"${camp_match.group(1)}/kit"
        shuttle_match = re.search(r'Shuttle.*?\$(\d+)', desc, re.IGNORECASE)
        if shuttle_match: shuttle_fee = f"${shuttle_match.group(1)}"

    if bike_fee != 'N/A' or camp_fee != 'N/A' or shuttle_fee != 'N/A':
        fees_section = f"""
## 💰 Fees & Logistics
| Item | Cost / Details |
| :--- | :--- |
| **Bike Rental** | {bike_fee} |
| **Camp Kit** | {camp_fee} |
| **Shuttle Service** | {shuttle_fee} |
"""

    return f"""# {title}

<!-- SYSTEM METADATA -->
| Arctic Code | System Status | Website ID |
| :--- | :--- | :--- |
| **{arctic_code}** | {tour.get('Sync_Status', 'Active')} | {tour.get('Website_ID', 'N/A')} |

---

## 1. The Shared DNA
**Subtitle:** {meta.get('Subtitle', '')}  
**Region:** {meta.get('Region', '')}  
**Skill Level:** {meta.get('Skill', '')}  
**Season:** {meta.get('Season', '')}

**Short Description:**
> {tour.get('Description_Short', '')}

**Long Description:**
> {tour.get('Description_Long', '')[:1500] if tour.get('Description_Long') else ''}...

**Images (Filenames):**
`{meta.get('Images', 'No images found')}`

---

## 2. Arctic Configurations (SKUs)
| Variant Name | Arctic ID | Price | Duration | Type |
| :--- | :--- | :--- | :--- | :--- |
{v_table}

{fees_section}

---

## 3. Itinerary Logic
| Day | Route Description | Miles | Elev Gain | Camp/Lodging | Applies To |
| :--- | :--- | :--- | :--- | :--- | :--- |
{i_table}

---

## 4. Full Itinerary Text
{full_text}
"""

def main():
    # Load Data
    try:
        with open(TOURS_FILE, 'r', encoding='utf-8') as f: tours = json.load(f)
    except: return print("Run unify_tours_advanced.py first!")

    # Load Itineraries from both directories
    doc_db_new = load_docx_itineraries("New Itineraries")
    doc_db_old = load_docx_itineraries("Old Itineraries")

    # Combine both document databases
    doc_db = {**doc_db_new, **doc_db_old}

    # Create Output Dir
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)

    for tour in tours:
        name = tour['Master_Name']
        if any(x in name.lower() for x in ['fake', 'test', 'sold out', 'aaa ']): continue

        # Find Itinerary Doc
        matched_doc_name = None
        itinerary_data = None

        # Fuzzy Search - use more nuanced matching
        best_score = 0
        matched_doc_name = None  # Reset in case it was set earlier

        for fname in doc_db:
            # Clean the filename for comparison, removing extensions and separators
            clean = fname.replace(".docx","").replace(".doc","").replace("-"," ").replace("_"," ")
            # Calculate similarity between tour name and clean filename
            score = SequenceMatcher(None, name.lower(), clean.lower()).ratio()

            # For The Maze specifically, allow lower threshold since it has multiple day variations
            if name.lower() == "the maze":
                required_threshold = 0.40  # Lower threshold for The Maze
            else:
                required_threshold = 0.65

            if score > required_threshold and score > best_score:
                best_score = score
                matched_doc_name = fname

        if matched_doc_name:
            itinerary_data = process_doc_content(doc_db[matched_doc_name])
            print(f"✅ Found Doc for {name}: {matched_doc_name} (parsed {len(itinerary_data)} days)")
        else:
            # For day tours, no itinerary is needed
            category = categorize(tour)
            if category == "Day Tours":
                itinerary_data = None
                print(f"ℹ️  Day Tour (no itinerary needed): {name}")
            else:
                print(f"⚠️  No Doc found for {name}")
                itinerary_data = None

        # Generate
        md = generate_markdown(tour, itinerary_data)

        # Save to proper subfolder
        sub = categorize(tour)
        p = os.path.join(OUTPUT_DIR, sub)
        if not os.path.exists(p): os.makedirs(p)
        safe = re.sub(r'[^\w\-_\. ]', '_', name) + ".md"
        with open(os.path.join(p, safe), 'w', encoding='utf-8') as f: f.write(md)

        print(f"📝 Generated {name}")

if __name__ == "__main__":
    main()