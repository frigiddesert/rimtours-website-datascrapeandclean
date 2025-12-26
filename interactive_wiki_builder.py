import json
import os
import re
import sys
import pandas as pd
import docx
from difflib import SequenceMatcher

# ==========================================
# CONFIGURATION
# ==========================================
OUTPUT_DIR = "outline_final_import"
PRICING_FILE = "arctic_pricing_final.csv"
TOURS_FILE = "unified_tours.json"
DOCS_FOLDER = "itinerary_docs"
REPORT_FILE = "build_report.csv"

# ==========================================
# 1. HELPERS & LOADERS
# ==========================================
def extract_stats_from_text(text):
    """Extracts Miles, Elev, Camp from a raw string."""
    if not text:
        return {"Miles": "", "Elev": "", "Camp": ""}
    
    stats = {"Miles": "", "Elev": "", "Camp": ""}
    # Miles
    m = re.search(r"(\d+(?:\.\d+)?)\s*(?:miles|mi\b)", text, re.IGNORECASE)
    if m: stats["Miles"] = m.group(1)
    # Elev
    e = re.search(r"(\d{1,3}(?:,\d{3})?)\s*(?:ft|feet|’|')\s*(?:gain|climb)", text, re.IGNORECASE)
    if e: stats["Elev"] = e.group(1)
    # Camp
    c = re.search(r"(?:Camp|Lodging|Overnight):\s*(.*)", text, re.IGNORECASE)
    if c: stats["Camp"] = c.group(1).strip()
    return stats

def load_pricing_map():
    try:
        df = pd.read_csv(PRICING_FILE)
    except:
        return {}
    price_map = {}
    grouped = df.groupby('Arctic_ID')
    for arctic_id, group in grouped:
        # Filter logic (same as before)
        valid = group[~group['Price_Name'].str.contains('Solo|solo|MAC|Viator|Net', case=False, na=False)]
        std = valid[valid['Price_Name'].str.contains('Standard|Adult|2\+', case=False, na=False)]
        target = std.iloc[0] if not std.empty else (valid.iloc[0] if not valid.empty else None)
        
        if target is not None:
            amt = float(target['Amount'])
            price_map[str(arctic_id)] = f"${amt:,.0f}" if amt < 200000 else "Call"
    return price_map

def load_docx_itineraries():
    db = {}
    if not os.path.exists(DOCS_FOLDER): return db
    
    files = [f for f in os.listdir(DOCS_FOLDER) if f.lower().endswith(('.doc', '.docx'))]
    print(f"📁 Found {len(files)} itinerary documents...")
    
    # For .docx files
    import subprocess
    for f in files:
        filepath = os.path.join(DOCS_FOLDER, f)
        try:
            if f.lower().endswith('.docx'):
                doc = docx.Document(filepath)
                db[f] = doc # Store doc ref to process later if matched
            elif f.lower().endswith('.doc'):
                # Use antiword to extract text
                result = subprocess.run(['antiword', filepath], 
                                      capture_output=True, check=True)
                try:
                    text = result.stdout.decode('utf-8', errors='ignore')
                except:
                    text = result.stdout.decode('latin-1', errors='ignore')
                # Create a mock object with paragraphs attribute for consistency
                class MockDoc:
                    def __init__(self, text):
                        self.paragraphs = [MockPara(line) for line in text.split('\n') if line.strip()]
                class MockPara:
                    def __init__(self, text):
                        self.text = text
                db[f] = MockDoc(text)
        except Exception as e:
            print(f"  Error reading {f}: {e}")
            continue
    return db

def process_doc_content(doc):
    """Parses a loaded Docx/Doc object into our list format."""
    days = []
    day_pattern = re.compile(r'(?:^|\n)\s*(?:Day|DAY|Stage|STAGE)\s+(\d+|[Oo]ne|[Tt]wo|[Tt]hree|[Ff]our|[Ff]ive|[Ss]ix)(?:[:\.-]|\s|$)', re.IGNORECASE)
    
    # Get all text first
    full_text = '\n'.join([para.text for para in doc.paragraphs if para.text])
    
    # Find all day markers
    matches = list(day_pattern.finditer(full_text))
    
    if not matches:
        # If no day markers, return as single day
        stats = extract_stats_from_text(full_text)
        return [{
            "Day": "1",
            "Content": full_text[:500],
            "Miles": stats['Miles'],
            "Elev": stats['Elev'],
            "Camp": stats['Camp']
        }]
    
    # Process each day segment
    mapping = {'one':'1','two':'2','three':'3','four':'4','five':'5','six':'6'}
    
    for i in range(len(matches)):
        match = matches[i]
        day_num = mapping.get(match.group(1).lower(), match.group(1))
        
        # Determine content boundaries
        start_pos = match.end()
        end_pos = matches[i+1].start() if i+1 < len(matches) else len(full_text)
        
        content = full_text[start_pos:end_pos].strip()
        
        # Extract stats from content
        stats = extract_stats_from_text(content)
        
        days.append({
            "Day": day_num,
            "Content": content,
            "Miles": stats['Miles'],
            "Elev": stats['Elev'],
            "Camp": stats['Camp']
        })
    
    return days

# ==========================================
# 2. INTERACTIVE MANUAL ENTRY (Updated for Environment Compatibility)
# ==========================================
def manual_entry_wizard(tour_name):
    # For automated environments, we'll skip manual entry and report it
    print(f"   → Skipping manual entry for {tour_name} (automated environment)")
    return None, "Manual Entry Needed"

# ==========================================
# 3. GENERATOR & MAIN
# ==========================================
def categorize(tour):
    name = tour['Master_Name'].lower()
    variants = tour['Arctic_Variants']
    if "oregon" in name or "sierra" in name: return "Archive"
    if "rental" in name or "service" in name or "shuttle" in name: return "Rentals & Services"
    is_multi = False
    for v in variants:
        # Updated to use the actual field name from the data
        bg = str(v.get('businessgroupid', ''))
        if bg in ['3','4','23','26']: is_multi = True
    if not variants and ("4-day" in name or "3-day" in name): is_multi = True
    
    if not is_multi: return "Day Tours"
    if "arizona" in name: return "Multi-Day Tours/Arizona"
    if "colorado" in name: return "Multi-Day Tours/Colorado"
    return "Multi-Day Tours/Utah"

def generate_md(tour, itinerary, price_map):
    # Build Pricing Table ...
    v_rows = []
    for v in tour['Arctic_Variants']:
        # Access the actual field names from the data
        aid = str(v.get('id', 'N/A'))
        v_name = v.get('name', 'N/A')
        v_duration = v.get('duration', 'N/A')
        v_type = v.get('name', 'N/A')  # or extract from keywords
        
        p = price_map.get(aid, "TBD")
        v_rows.append(f"| **{v_name}** | {aid} | {p} | {v_duration} | {v_type} |")
    v_table = "\n".join(v_rows) if v_rows else "| No Linked Variants | - | - | - | - |"

    # Build Itinerary Table ...
    i_rows = []
    full_txt = ""
    if itinerary:
        for d in itinerary:
            desc = d['Content'][:100].replace("\n"," ") + "..."
            i_rows.append(f"| {d['Day']} | {desc} | {d['Miles']} | {d['Elev']} | {d['Camp']} | All |")
            full_txt += f"### Day {d['Day']}\n{d['Content']}\n\n"
    else:
        i_rows = ["| 1 | | | | | All |"]
    i_table = "\n".join(i_rows)

    return f"""# {tour['Master_Name']}

<!-- SYSTEM METADATA -->
| System | Status | Master ID |
| :--- | :--- | :--- |
| **SSOT** | Active | {tour.get('Website_ID','New')} |
| **Web** | {tour.get('Sync_Status','Linked')} | {tour.get('Slug','N/A')} |

---
## 1. Overview
> {tour.get('Description_Short','')}

---
## 2. Arctic Configurations
| Variant Name | Arctic ID | Price | Duration | Type |
| :--- | :--- | :--- | :--- | :--- |
{v_table}

---
## 3. Itinerary Logic
| Day | Route | Miles | Elev | Camp | Scope |
| :--- | :--- | :--- | :--- | :--- | :--- |
{i_table}

---
## 4. Full Itinerary
{full_txt}
"""

def main():
    try:
        with open(TOURS_FILE, 'r') as f: tours = json.load(f)
    except: return print("Missing unified_tours.json")

    prices = load_pricing_map()
    doc_db = load_docx_itineraries()
    
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)
    
    report_data = []

    print(f"--- Processing {len(tours)} Tours ---")

    for tour in tours:
        name = tour['Master_Name']
        if any(x in name.lower() for x in ['fake', 'test', 'sold out']): continue
        
        # 1. Try Find Itinerary Doc
        matched_doc_name = None
        itinerary_data = None
        source_status = "None"
        
        # Fuzzy Search - use both .doc and .docx
        best_score = 0
        for fname in doc_db:
            clean = fname.replace(".docx","").replace(".doc","").replace("-"," ").replace("_"," ")
            score = SequenceMatcher(None, name.lower(), clean.lower()).ratio()
            if score > 0.65 and score > best_score:
                best_score = score
                matched_doc_name = fname
        
        if matched_doc_name:
            itinerary_data = process_doc_content(doc_db[matched_doc_name])
            source_status = "Doc"  # Updated to reflect both .doc and .docx
            print(f"✅ Found Doc for {name}: {matched_doc_name}")
        else:
            # 2. Manual Fallback - For day tours, we won't ask for manual entry as mentioned
            category = categorize(tour)
            if category == "Day Tours":
                # Day tours don't need detailed itineraries, just skip
                itinerary_data, source_status = None, "No Itinerary Needed"
                print(f"ℹ️  Day Tour (no itinerary needed): {name}")
            elif category == "Archive":
                # Archive tours don't need detailed itineraries, just skip
                itinerary_data, source_status = None, "Archive - No Itinerary Needed"
                print(f"📦 Archive Tour (no itinerary needed): {name}")
            elif category == "Rentals & Services":
                # Rentals/Services don't need detailed itineraries, just skip
                itinerary_data, source_status = None, "Rental/Service - No Itinerary Needed"
                print(f"🛠️  Rental/Service (no itinerary needed): {name}")
            else:
                # For multi-day tours without docs, we'll note that manual entry was needed
                print(f"⚠️  No Doc found for {name}")
                itinerary_data, source_status = manual_entry_wizard(name)

        # 3. Generate
        md = generate_md(tour, itinerary_data, prices)
        
        # 4. Save
        sub = categorize(tour)
        p = os.path.join(OUTPUT_DIR, sub)
        if not os.path.exists(p): os.makedirs(p)
        safe = re.sub(r'[^\w\-_\. ]', '_', name) + ".md"
        with open(os.path.join(p, safe), 'w', encoding='utf-8') as f: f.write(md)

        # 5. Log Report
        report_data.append({
            "Tour Name": name,
            "Category": sub,
            "Itinerary Source": source_status,
            "Price Points": len(tour['Arctic_Variants']),
            "Filename": safe,
            "Days Parsed": len(itinerary_data) if itinerary_data else 0
        })

    # Save Report
    pd.DataFrame(report_data).to_csv(REPORT_FILE, index=False)
    print(f"\nReport saved to {REPORT_FILE}")
    print(f"\nCompleted! All tours organized in '{OUTPUT_DIR}/' folder.")
    print("For tours that need manual entry, use the interactive version in a terminal environment.")

if __name__ == "__main__":
    main()