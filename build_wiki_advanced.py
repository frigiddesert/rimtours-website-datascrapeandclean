import json
import os
import re
import pandas as pd
import docx
from difflib import SequenceMatcher

# CONFIGURATION
OUTPUT_DIR = "outline_final_import"
TOURS_FILE = "unified_tours.json"
DOCS_FOLDER = "itinerary_docs"

# ==========================================
# 1. HELPER FUNCTIONS
# ==========================================

def extract_stats_from_text(text):
    """Extracts Miles, Elev, Camp from a raw string."""
    if not text:
        return {"Miles": "", "Elev": "", "Camp": ""}
    
    stats = {"Miles": "", "Elev": "", "Camp": ""}
    # Miles
    m = re.search(r"(\d+(?:\.\d+)?)\s*(?:miles|mi\b)", text, re.IGNORECASE)
    if m: stats["Miles"] = m.group(1)
    # Elev
    e = re.search(r"(\d{1,3}(?:,\d{3})?)\s*(?:ft|feet|’|')\s*(?:gain|climb)", text, re.IGNORECASE)
    if e: stats["Elev"] = e.group(1)
    # Camp
    c = re.search(r"(?:Camp|Lodging|Overnight):\s*(.*)", text, re.IGNORECASE)
    if c: stats["Camp"] = c.group(1).strip()
    return stats

def load_docx_itineraries():
    db = {}
    if not os.path.exists(DOCS_FOLDER): 
        return db
    
    files = [f for f in os.listdir(DOCS_FOLDER) if f.lower().endswith(('.doc', '.docx'))]
    print(f"📁 Found {len(files)} itinerary documents...")
    
    # For .docx files
    import subprocess
    for f in files:
        filepath = os.path.join(DOCS_FOLDER, f)
        try:
            if f.lower().endswith('.docx'):
                doc = docx.Document(filepath)
                db[f] = doc # Store doc ref to process later if matched
            elif f.lower().endswith('.doc'):
                # Use antiword to extract text
                result = subprocess.run(['antiword', filepath], 
                                      capture_output=True, check=True)
                try:
                    text = result.stdout.decode('utf-8', errors='ignore')
                except:
                    text = result.stdout.decode('latin-1', errors='ignore')
                # Create a mock object with paragraphs attribute for consistency
                class MockDoc:
                    def __init__(self, text):
                        self.paragraphs = [MockPara(line) for line in text.split('\n') if line.strip()]
                class MockPara:
                    def __init__(self, text):
                        self.text = text
                db[f] = MockDoc(text)
        except Exception as e:
            print(f"  Error reading {f}: {e}")
            continue
    return db

def process_doc_content(doc):
    """Parses a loaded Docx/Doc object into our list format."""
    days = []
    day_pattern = re.compile(r'(?:^|\n)\s*(?:Day|DAY|Stage|STAGE)\s+(\d+|[Oo]ne|[Tt]wo|[Tt]hree|[Ff]our|[Ff]ive|[Ss]ix)(?:[:\.-]|\s|$)', re.IGNORECASE)
    
    # Get all text first
    full_text = '\n'.join([para.text for para in doc.paragraphs if para.text])
    
    # Find all day markers
    matches = list(day_pattern.finditer(full_text))
    
    if not matches:
        # If no day markers, return as single day
        stats = extract_stats_from_text(full_text)
        return [{
            "Day": "1",
            "Content": full_text[:500],
            "Miles": stats['Miles'],
            "Elev": stats['Elev'],
            "Camp": stats['Camp']
        }]
    
    # Process each day segment
    mapping = {'one':'1','two':'2','three':'3','four':'4','five':'5','six':'6'}
    
    for i in range(len(matches)):
        match = matches[i]
        day_num = mapping.get(match.group(1).lower(), match.group(1))
        
        # Determine content boundaries
        start_pos = match.end()
        end_pos = matches[i+1].start() if i+1 < len(matches) else len(full_text)
        
        content = full_text[start_pos:end_pos].strip()
        
        # Extract stats from content
        stats = extract_stats_from_text(content)
        
        days.append({
            "Day": day_num,
            "Content": content,
            "Miles": stats['Miles'],
            "Elev": stats['Elev'],
            "Camp": stats['Camp']
        })
    
    return days


def categorize(tour):
    name = tour['Master_Name'].lower()
    variants = tour['Arctic_Variants']
    if "oregon" in name or "sierra" in name: return "Archive"
    if "rental" in name or "service" in name or "shuttle" in name: return "Rentals & Services"
    is_multi = False
    for v in variants:
        # Updated to use the actual field name from the data
        bg = str(v.get('Business_Group', ''))
        if bg in ['3','4','23','26']: is_multi = True
    if not variants and ("4-day" in name or "3-day" in name): is_multi = True
    
    if not is_multi: return "Day Tours"
    if "arizona" in name: return "Multi-Day Tours/Arizona"
    if "colorado" in name: return "Multi-Day Tours/Colorado"
    return "Multi-Day Tours/Utah"


def generate_markdown(tour, itinerary):
    title = tour['Master_Name']
    meta = tour.get('Meta', {})
    
    # 1. Get Primary Arctic Shortname (from the first variant, or list them)
    shortnames = sorted(list(set([v['Shortname'] for v in tour['Arctic_Variants'] if v.get('Shortname')])))
    arctic_code = ", ".join(shortnames) if shortnames else "N/A"

    # 2. Variants Table
    v_rows = []
    for v in tour['Arctic_Variants']:
        # | Shortname | Variant Name | ID | Price | Duration |
        row = f"| {v.get('Shortname','')} | **{v['Name']}** | {v['Arctic_ID']} | {v['Price']} | {v['Duration']} |"
        v_rows.append(row)
    v_table = "\n".join(v_rows) if v_rows else "| - | No Linked Variants | - | - | - |"

    # 3. Itinerary Table
    i_rows = []
    full_text = ""
    if itinerary:
        for d in itinerary:
            desc = d['Content'][:150].replace("\n", " ") + "..."
            i_rows.append(f"| {d['Day']} | {desc} | {d['Miles']} | {d['Elev']} | {d['Camp']} | All |")
            full_text += f"### Day {d['Day']}\n{d['Content']}\n\n"
    else:
        i_rows = ["| 1 | | | | | All |", "| 2 | | | | | All |"]
    i_table = "\n".join(i_rows)

    # 4. Fees Block
    fees_section = ""
    if meta.get('Fees_Bike') or meta.get('Fees_Camp') or meta.get('Fees_Shuttle'):
        fees_section = f"""
## 💰 Fees & Add-ons
| Item | Cost / Details |
| :--- | :--- |
| **Bike Rental** | {meta.get('Fees_Bike', 'N/A')} |
| **Camp Kit** | {meta.get('Fees_Camp', 'N/A')} |
| **Shuttle** | {meta.get('Fees_Shuttle', 'Included or N/A')} |
"""

    return f"""# {title}

<!-- SYSTEM METADATA -->
| Arctic Code | System Status | Website ID |
| :--- | :--- | :--- |
| **{arctic_code}** | {tour.get('Sync_Status')} | {tour.get('Website_ID', 'N/A')} |

---

## 1. The Shared DNA
**Subtitle:** {meta.get('Subtitle', '')}  
**Region:** {meta.get('Region', '')}  
**Skill Level:** {meta.get('Skill', '')}  
**Season:** {meta.get('Season', '')}

**Short Description:**
> {tour.get('Description_Short', '')}

**Long Description:**
> {tour.get('Description_Long', '')[:1500]}...

**Images (Filenames):**
`{meta.get('Images', 'No images found')}`

---

## 2. Arctic Configurations (SKUs)
| Code | Variant Name | Arctic ID | Price | Duration |
| :--- | :--- | :--- | :--- | :--- |
{v_table}

{fees_section}

---

## 3. Itinerary Logic
| Day | Route Description | Miles | Elev Gain | Camp/Lodging | Applies To |
| :--- | :--- | :--- | :--- | :--- | :--- |
{i_table}

---

## 4. Full Itinerary Text
{full_text}
"""

def main():
    # Load Data
    try:
        with open(TOURS_FILE, 'r') as f: tours = json.load(f)
    except: return print("Run unify_tours.py first!")
    
    # Load Itineraries
    doc_db = load_docx_itineraries()
    
    # Create Output
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)
    
    for tour in tours:
        name = tour['Master_Name']
        if any(x in name.lower() for x in ['fake', 'test', 'sold out', 'aaa']): continue
        
        # 1. Try Find Itinerary Doc
        matched_doc_name = None
        itinerary_data = None
        
        # Fuzzy Search - use both .doc and .docx
        best_score = 0
        for fname in doc_db:
            clean = fname.replace(".docx","").replace(".doc","").replace("-"," ").replace("_"," ")
            score = SequenceMatcher(None, name.lower(), clean.lower()).ratio()
            if score > 0.65 and score > best_score:
                best_score = score
                matched_doc_name = fname
        
        if matched_doc_name:
            itinerary_data = process_doc_content(doc_db[matched_doc_name])
            print(f"✅ Found Doc for {name}: {matched_doc_name}")
        else:
            # For day tours, no itinerary is needed, just skip
            category = categorize(tour)
            if category == "Day Tours":
                # Day tours don't need detailed itineraries, just skip
                itinerary_data = None
                print(f"ℹ️  Day Tour (no itinerary needed): {name}")
            else:
                # Only for multi-day tours that don't have docs
                print(f"⚠️  No Doc found for {name}")
                itinerary_data = None

        # Categorize tour
        subfolder = categorize(tour)
        
        # Generate markdown
        md = generate_markdown(tour, itinerary_data)
        
        # Save to correct subfolder
        p = os.path.join(OUTPUT_DIR, subfolder)
        if not os.path.exists(p): os.makedirs(p)
        safe = re.sub(r'[^\w\-_\. ]', '_', name) + ".md"
        with open(os.path.join(p, safe), 'w', encoding='utf-8') as f: f.write(md)

        print(f"📝 Generated {name}")

if __name__ == "__main__":
    main()